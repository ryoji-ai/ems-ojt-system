"""
generate-annual-plan.py

annual-plan-2026.json を読み込み、決裁用 Word 文書を生成するスクリプト。

実行方法:
    pip install python-docx
    python docs/generate-annual-plan.py
    → docs/output/年間訓練計画書_令和8年度.docx が生成される
"""

import json
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ── ファイルパス設定 ──────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PLAN_FILE = os.path.join(SCRIPT_DIR, 'annual-plan-2026.json')
OUTPUT_DIR = os.path.join(SCRIPT_DIR, 'output')
OUTPUT_FILE = os.path.join(OUTPUT_DIR, '年間訓練計画書_令和8年度.docx')

os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── 月名変換テーブル ─────────────────────────────────────────────────────
MONTH_LABELS = {
    '2026-04': '4月（令和8年）',
    '2026-05': '5月（令和8年）',
    '2026-06': '6月（令和8年）',
    '2026-07': '7月（令和8年）',
    '2026-08': '8月（令和8年）',
    '2026-09': '9月（令和8年）',
    '2026-10': '10月（令和8年）',
    '2026-11': '11月（令和8年）',
    '2026-12': '12月（令和8年）',
    '2027-01': '1月（令和9年）',
    '2027-02': '2月（令和9年）',
    '2027-03': '3月（令和9年）',
}

# ── ヘルパー関数 ─────────────────────────────────────────────────────────
def set_run_font(run, font_name='游明朝', size=10.5):
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_section_heading(doc, text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size)
    run.bold = True
    return p

def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(('　' if indent else '') + text)
    set_run_font(run)
    return p

# ── JSON読み込み ─────────────────────────────────────────────────────────
with open(PLAN_FILE, encoding='utf-8') as f:
    plan = json.load(f)

months = plan['months']

# ── 文書生成 ─────────────────────────────────────────────────────────────
doc = Document()

# ── 1. タイトル ──────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('起　案　文　書')
set_run_font(run, size=16)
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('件　名：')
set_run_font(run)
run.bold = True
run = p.add_run(plan['title'] + 'について')
set_run_font(run)

p = doc.add_paragraph()
run = p.add_run('起案日：')
set_run_font(run)
run.bold = True
run = p.add_run('令和8年4月1日')
set_run_font(run)

p = doc.add_paragraph()
run = p.add_run('起案者：')
set_run_font(run)
run.bold = True
run = p.add_run(plan['organization'] + ' ' + plan['author'] + '　指導救命士　○○○○')
set_run_font(run)

doc.add_paragraph()

# ── 2. 趣旨 ─────────────────────────────────────────────────────────────
add_section_heading(doc, '1. 趣　旨')
add_body(doc, '　令和8年度の救急訓練について、下記のとおり年間計画を策定し、計画的・効率的な救急隊員教育を実施したい。')
add_body(doc, '　本計画を年度初めに一括決裁することで、毎月の訓練起案を不要とし、指導救命士の事務負担を軽減する。')

doc.add_paragraph()

# ── 3. 訓練計画の概要 ───────────────────────────────────────────────────
add_section_heading(doc, '2. 訓練計画の概要')

table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
overview_data = [
    ('項目', '内容'),
    ('対象年度', '令和8年度（2026年4月〜2027年3月）'),
    ('実施頻度', '月1回（毎月の訓練目標に基づき実施）'),
    ('対象者', '教育担当隊長 8名（本署4名、西出張所2名、昭和出張所2名）'),
    ('記録方法', '訓練実施後、システム内「記録する」ボタンからGoogle Formに入力'),
]
for i, (col1, col2) in enumerate(overview_data):
    run1 = table.rows[i].cells[0].paragraphs[0].add_run(col1)
    set_run_font(run1)
    if i == 0:
        run1.bold = True
    run2 = table.rows[i].cells[1].paragraphs[0].add_run(col2)
    set_run_font(run2)
    if i == 0:
        run2.bold = True

doc.add_paragraph()

# ── 4. 年間訓練計画表 ───────────────────────────────────────────────────
add_section_heading(doc, '3. 年間訓練計画表')

# ヘッダー行 + 12ヶ月 = 13行, 4列（月・テーマ・訓練目標・評価基準）
num_rows = 1 + len(months)
plan_table = doc.add_table(rows=num_rows, cols=4)
plan_table.style = 'Table Grid'

# ヘッダー
headers = ['月', 'テーマ', '訓練目標', '主な評価基準']
for i, h in enumerate(headers):
    run = plan_table.rows[0].cells[i].paragraphs[0].add_run(h)
    set_run_font(run)
    run.bold = True
    plan_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 各月のデータ
for row_idx, (month_key, data) in enumerate(months.items(), start=1):
    row = plan_table.rows[row_idx]

    # 月
    run = row.cells[0].paragraphs[0].add_run(MONTH_LABELS.get(month_key, month_key))
    set_run_font(run, size=9)

    # テーマ
    run = row.cells[1].paragraphs[0].add_run(data['theme'])
    set_run_font(run, size=9)

    # 訓練目標（複数）
    for g in data['goals']:
        p = row.cells[2].add_paragraph()
        run = p.add_run('・' + g['title'])
        set_run_font(run, size=9)
        run.bold = True
        p2 = row.cells[2].add_paragraph()
        run2 = p2.add_run('　' + g['description'])
        set_run_font(run2, size=8.5)

    # 評価基準（複数）
    for ev in data['evaluation']:
        p = row.cells[3].add_paragraph()
        run = p.add_run('□ ' + ev)
        set_run_font(run, size=8.5)

doc.add_paragraph()

# ── 5. 実施要領 ─────────────────────────────────────────────────────────
add_section_heading(doc, '4. 実施要領')
items = [
    ('訓練形式', '実演型（実際に手技を行わせて確認する形式）を基本とする'),
    ('所要時間', '1回あたり30〜60分程度'),
    ('実施場所', '各署・出張所の訓練室または車庫内'),
    ('シナリオ', '本計画の各月シナリオを参考に、当日の訓練状況に応じて柔軟に対応する'),
    ('記録', '訓練実施後はOJT教育システムの記録機能（Google Form）を活用して記録を残す'),
    ('フィードバック', '訓練終了後、指導担当者から評価基準に基づいたフィードバックを実施する'),
]

fb_table = doc.add_table(rows=len(items), cols=2)
fb_table.style = 'Table Grid'
for i, (item, content) in enumerate(items):
    run1 = fb_table.rows[i].cells[0].paragraphs[0].add_run(item)
    set_run_font(run1)
    run1.bold = True
    run2 = fb_table.rows[i].cells[1].paragraphs[0].add_run(content)
    set_run_font(run2)

doc.add_paragraph()

# ── 6. その他 ───────────────────────────────────────────────────────────
add_section_heading(doc, '5. その他')
others = [
    '岡山県MCプロトコル改訂時は、関連月の訓練内容を適宜修正する',
    '年度末（3月）に訓練実施状況を振り返り、翌年度計画の参考とする',
    '訓練内容の更新は OJT教育システム（monthly-goals.js）に反映し、GitHub Pagesに公開する',
]
for i, item in enumerate(others, 1):
    add_body(doc, f'({i}) {item}')

doc.add_paragraph()
doc.add_paragraph()

# ── 結び ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('上記のとおり起案する。')
set_run_font(run)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

# ── 決裁欄 ──────────────────────────────────────────────────────────────
approval_table = doc.add_table(rows=2, cols=5)
approval_table.style = 'Table Grid'
approvers = ['消防長', '次長', '救急課長', '係長', '起案者']
for i, name in enumerate(approvers):
    cell = approval_table.rows[0].cells[i]
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell.paragraphs[0].add_run(name)
    set_run_font(run, size=9)
for i in range(5):
    approval_table.rows[1].cells[i].paragraphs[0].add_run('\n\n\n')

# ── 保存 ────────────────────────────────────────────────────────────────
doc.save(OUTPUT_FILE)
print(f'Word文書を生成しました: {OUTPUT_FILE}')
